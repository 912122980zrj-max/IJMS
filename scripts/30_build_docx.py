#!/usr/bin/env python
"""30_build_docx.py

把 pandoc 生成的 DOCX 后处理为 PLOS 兼容稿件：
  - 正文字体 Times New Roman 12 pt；
  - 正文双倍行距（表格单倍行距）；
  - 连续行号（countBy=1, restart=continuous）；
  - 页脚居中页码（PAGE 字段）；
  - 输出时使用新文件名，不覆盖原始 pandoc 产物。
"""

from __future__ import annotations

import argparse
import copy
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_margins(table):
    tbl_pr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for side, w in (("top", 57), ("start", 108), ("bottom", 57), ("end", 108)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(w))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)


def add_line_numbering(doc):
    # 在 settings.xml 与每个 section 的 sectPr 中启用连续行号
    settings = doc.settings.element
    if settings.find(qn("w:lnNumType")) is None:
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:restart"), "continuous")
        settings.append(ln)
    for section in doc.sections:
        sect_pr = section._sectPr
        old = sect_pr.find(qn("w:lnNumType"))
        if old is not None:
            sect_pr.remove(old)
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:restart"), "continuous")
        sect_pr.append(ln)


def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run.font.name = "Times New Roman"
    run.font.size = None


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = None
    normal.paragraph_format.space_after = None
    rpr = normal.element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "Times New Roman")
    fonts.set(qn("w:cs"), "Times New Roman")
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            continue
        fmt = p.paragraph_format
        fmt.line_spacing = 2.0
        fmt.space_before = None
        fmt.space_after = None
    for table in doc.tables:
        set_cell_margins(table)
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.0
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = None


def main() -> int:
    ap = argparse.ArgumentParser(description="PLOS DOCX 后处理")
    ap.add_argument("input")
    ap.add_argument("output")
    ap.add_argument("--force", action="store_true", help="允许覆盖输出")
    args = ap.parse_args()
    if not os.path.exists(args.input):
        raise SystemExit(f"missing input: {args.input}")
    if os.path.exists(args.output) and not args.force:
        raise SystemExit(f"refusing to overwrite: {args.output}")
    doc = Document(args.input)
    style_document(doc)
    add_line_numbering(doc)
    add_page_number_footer(doc)
    doc.save(args.output)
    print(f"written -> {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
