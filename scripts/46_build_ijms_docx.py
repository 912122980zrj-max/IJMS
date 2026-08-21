#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""46_build_ijms_docx.py —— 将 IJMS 稿件 Markdown 转为带插图留白区的 Word 稿件

输入: submission/ijms/manuscript/manuscript_ijms_draft.md
输出: submission/ijms/manuscript/manuscript_ijms_draft.docx

要点:
  - 正文 Palatino Linotype 10 pt（MDPI 模板惯用字体，Word 可自动回退）；
  - 每个 Figure 在首次引用段落后插入一个带边框的空白占位框（可编辑表单元格），
    框内仅写提示文字，高度 8–10 cm，供后续插入 300 dpi 图件；
  - 图注置于占位框下方（MDPI：图注在图下方），表题置于表上方；
  - 把 "× 10−n"、"2−ΔΔCt" 渲染为上标；解析 **加粗** 与 *斜体* 行内标记；
  - 不覆盖已有文件（需 --force）。
"""

from __future__ import annotations

import argparse
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

ROOT = r"E:/sheng xin/ObstructiveNephropathy_MRG"
MD = os.path.join(ROOT, "submission", "ijms", "manuscript",
                  "manuscript_ijms_draft.md")
OUT = os.path.join(ROOT, "submission", "ijms", "manuscript",
                   "manuscript_ijms_draft.docx")
FIGDIR = os.path.join(ROOT, "submission", "ijms", "figures", "preview")

FONT = "Palatino Linotype"


def add_runs(par, text: str, base_bold: bool = False,
             base_italic: bool = False, size: float = 10.0):
    """解析行内 **bold** 与 *italic* 并写 runs；处理上标。"""
    # 先把已知科学计数法与 ΔΔCt 标为上标：用占位符避免与 ** 冲突
    supers = []

    def stash(m):
        supers.append(m.group(0))
        return f"\x01{len(supers) - 1}\x02"

    text = re.sub(r"10−\d+", stash, text)
    text = re.sub(r"2−ΔΔC[Tt]", stash, text)

    token_re = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")
    for tok in token_re.split(text):
        if not tok:
            continue
        bold = base_bold
        italic = base_italic
        if tok.startswith("**") and tok.endswith("**"):
            tok = tok[2:-2]
            bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            tok = tok[1:-1]
            italic = True
        # 拆分上标占位符
        parts = re.split(r"(\x01\d+\x02)", tok)
        for part in parts:
            if not part:
                continue
            m = re.fullmatch(r"\x01(\d+)\x02", part)
            if m:
                run = par.add_run(supers[int(m.group(1))])
                run.font.superscript = True
            else:
                run = par.add_run(part)
            run.bold = bold
            run.italic = italic
            run.font.name = FONT
            run.font.size = Pt(size)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), FONT)
            rFonts.set(qn("w:hAnsi"), FONT)
            rFonts.set(qn("w:eastAsia"), FONT)


def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.2)


def heading(doc, text: str, level: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    if level == 1:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = FONT
    else:
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(10.5)
        run.font.name = FONT
    return p


def set_cell_shading(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_row_height(row, cm: float):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(cm * 567)))  # 1 cm = 567 twips
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def insert_figure(doc, fig_no: int, width_cm: float = 17.8, max_h_cm: float = 13.5):
    """插入实际图件（居中），限制最长边以免超高。找不到图时回退为留白提示框。"""
    img = os.path.join(FIGDIR, f"Fig{fig_no}.png")
    if not os.path.exists(img):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[Figure {fig_no} image not found — insert here]")
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        r.font.size = Pt(9)
        r.font.name = FONT
        return
    try:
        from PIL import Image
        im = Image.open(img)
        iw, ih = im.size
        # 依据宽高比计算插入尺寸：默认占满 17.8cm 宽，但高度不超过 max_h_cm
        w = width_cm
        h = width_cm * (ih / iw)
        if h > max_h_cm:
            scale = max_h_cm / h
            w *= scale
            h *= scale
    except Exception:
        w, h = width_cm, max_h_cm
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(img, width=Cm(w))
    return


def placeholder_box(doc, fig_no: int, text: str, height_cm: float):
    """居中、带边框的留白框 + 框内提示文字；现改为直接插入图件。"""
    insert_figure(doc, fig_no)
    return


def add_table_caption(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    m = re.match(r"^(\*\*Table \d+\.\*\*)\s*(.*)$", text)
    if m:
        add_runs(p, m.group(1), base_bold=True)
        p.add_run(" ")
        add_runs(p, m.group(2))
    else:
        add_runs(p, text)
    return p


def add_md_table(doc, rows: list[list[str]]):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], cell_text, base_bold=(i == 0),
                     size=9)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--force", action="store_true")
    args = ap.parse_args()
    if os.path.exists(OUT) and not args.force:
        raise SystemExit(f"refusing to overwrite: {OUT} (use --force)")

    lines = open(MD, encoding="utf-8").read().splitlines()
    doc = Document()
    style_base(doc)

    i = 0
    n_fig = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(17)
            run.font.name = FONT
            i += 1
            continue

        if stripped.startswith("## "):
            heading(doc, stripped[3:], 1)
            i += 1
            continue

        if stripped.startswith("### "):
            heading(doc, stripped[4:], 2)
            i += 1
            continue

        if stripped.startswith("---"):
            i += 1
            continue

        # 图占位
        mfig = re.match(r"^\*\*Figure (\d+) placeholder\*\*", stripped)
        if mfig:
            n_fig += 1
            no = int(mfig.group(1))
            height = 8.5 if no <= 3 else 10.5
            placeholder_box(
                doc, no,
                f"[Blank space — insert Figure {no} here at double-column "
                f"width (17.8 cm); height reserved {height:.1f} cm]",
                height)
            i += 1
            continue

        # 图注
        mcap = re.match(r"^\*\*Figure \d+\.\*\*", stripped)
        if mcap:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(10)
            add_runs(p, mcap.group(0), base_bold=True)
            p.add_run(" ")
            add_runs(p, stripped[mcap.end():].strip())
            i += 1
            continue

        # 表题
        if re.match(r"^\*\*Table \d+\.\*\*", stripped):
            add_table_caption(doc, stripped)
            i += 1
            continue

        # Markdown 表
        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            add_md_table(doc, rows)
            continue

        # 引用
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            add_runs(p, stripped, size=9.5)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, stripped)
        i += 1

    doc.save(OUT)
    print(f"[docx] written {OUT}")
    print(f"[docx] figure placeholders: {n_fig}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
