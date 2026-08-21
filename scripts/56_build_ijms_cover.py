#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""56_build_ijms_cover.py —— IJMS 投稿信 md → docx"""

from docx import Document
from docx.shared import Pt

OUT = (r"E:/sheng xin/ObstructiveNephropathy_MRG/submission/ijms/"
       "cover_letter/cover_letter_ijms.docx")

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Palatino Linotype"
st.font.size = Pt(10)

paras = [
    ("", "Date: 2026-08-21"),
    ("", ""),
    ("", "To the Editors"),
    ("italic", "International Journal of Molecular Sciences (IJMS)"),
    ("", ""),
    ("", "Dear Editors,"),
    ("", ""),
    ("", 'We are pleased to submit our manuscript entitled "A Mechanical '
          'Stress-Related Gene Signature Dissects Molecular Subtypes and '
          'Fibrotic Trajectories in Obstructive Nephropathy" by Renjie Zuo, '
          "Chenchun Ding and Cheng Sun (corresponding author) for "
          "consideration as an original research article in IJMS."),
    ("", ""),
    ("", "Obstructive nephropathy exposes the kidney to sustained hydrostatic "
          "pressure and tubular stretch, yet the transcriptomic organization "
          "of this mechanical stress response is poorly defined. Using public "
          "human kidney transcriptomes, mouse unilateral ureteral obstruction "
          "models, human spatial transcriptomics and pilot cell-culture "
          "experiments, we define a transcriptome-based mechanical axis: a "
          "three-gene signature (NDNF, PCDHB7, RRAGB) that discriminates "
          "immunoglobulin A nephropathy, two molecular subtypes with distinct "
          "pathway and immune activity, the collecting duct as a candidate "
          "mechanosensory hub, and disease-associated mechanical genes "
          "enriched for actin regulation and Hippo signaling. The analyses "
          "are fully scripted and reproducible, and the results provide a "
          "testable framework for mechanotransduction-oriented studies of "
          "renal fibrosis, fitting the scope of IJMS."),
    ("", ""),
    ("", "We confirm that neither the manuscript nor any parts of its content "
          "are currently under consideration for publication with or "
          "published in another journal."),
    ("", ""),
    ("", "All authors have approved the manuscript and agree with its "
          "submission to IJMS."),
    ("", ""),
    ("", "Thank you for your consideration."),
    ("", ""),
    ("", "Sincerely,"),
    ("", ""),
    ("", "Cheng Sun"),
    ("", "Department of Urology, The Third Affiliated Hospital of Anhui "
          "Medical University (The First People's Hospital of Hefei), Hefei, "
          "230061, China"),
    ("", "suncheng216@163.com"),
]

for style, text in paras:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Palatino Linotype"
    r.font.size = Pt(10)
    if style == "italic":
        r.italic = True

doc.save(OUT)
print("written", OUT)
