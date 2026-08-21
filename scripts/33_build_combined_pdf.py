#!/usr/bin/env python
"""33_build_combined_pdf.py

把 Fig1–4.tif 嵌入手稿 DOCX（紧跟各图题之后），另存为带图合并稿，
供 Word COM 导出单文件 PDF 用于初投稿审阅。不修改原始 manuscript_plos.docx。
"""

from __future__ import annotations

import argparse
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def main() -> int:
    ap = argparse.ArgumentParser(description="合并图稿 DOCX")
    ap.add_argument("--manuscript",
                    default=r"E:/sheng xin/ObstructiveNephropathy_MRG/submission/plos/manuscript/manuscript_plos.docx")
    ap.add_argument("--figdir",
                    default=r"E:/sheng xin/ObstructiveNephropathy_MRG/submission/plos/figures")
    ap.add_argument("--out",
                    default=r"E:/sheng xin/ObstructiveNephropathy_MRG/submission/plos/manuscript/manuscript_plos_with_figures.docx")
    ap.add_argument("--force", action="store_true", help="允许覆盖输出")
    args = ap.parse_args()

    if os.path.exists(args.out) and not args.force:
        raise SystemExit(f"refusing to overwrite: {args.out}")
    doc = Document(args.manuscript)
    legends = [p for p in doc.paragraphs
               if p.text.strip().startswith("Fig ")]
    if len(legends) != 4:
        raise SystemExit(f"expected 4 figure legends, found {len(legends)}")
    for i, para in enumerate(legends, start=1):
        tif = os.path.join(args.figdir, f"Fig{i}.tif")
        if not os.path.exists(tif):
            raise SystemExit(f"missing {tif}")
        # 在题注段后插入图片段
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic.add_run()
        run.add_picture(tif, width=Inches(6.5))
        # 把图片段移动到题注段之后
        para._p.addnext(pic._p)
    doc.save(args.out)
    print(f"written -> {args.out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
