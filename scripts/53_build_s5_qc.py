#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""53_build_s5_qc.py —— 生成 Text S5（qPCR 质控说明）docx"""

from docx import Document

OUT = (r"E:/sheng xin/ObstructiveNephropathy_MRG/submission/ijms/"
       "supporting_information/S5_Text_qpcr_qc.docx")

doc = Document()
doc.add_heading("Text S5. qPCR quality-control statement", level=1)

paras = [
    ("Plate 1 (mIMCD-3 mechanism plate) was recovered from a QuantStudio "
     "export. The instrument flagged SPIKE in 62 wells and HIGHSD in 49 "
     "wells. Melting-curve peak detection was 9/12 wells for each target "
     "gene, and several targets showed intra-target Tm ranges exceeding "
     "0.5 °C. In addition, the amplification-curve table could not be "
     "aligned to the reported Ct values (median offset of approximately "
     "−19.3 cycles), so the recovery attempt was abandoned and the "
     "compiled Ct values were used. Thirty-six wells carried manually "
     "entered Ct values with three to four decimal places; this rounding "
     "introduces a negligible effect (≤0.0005 cycles) on ΔΔCt."),
    ("Plates 2 and 3 are compiled tables without raw instrument exports; "
     "no machine QC flags were available for these plates."),
    ("Suspicious wells (A7 Piezo1, B11 Yap1, D9 Cyr61, E7 Ankrd1) are "
     "disclosed but were not removed."),
    ("All three plates are single pilot experiments with three technical "
     "replicate wells per group; mean fold changes and the range across "
     "the technical replicates are reported descriptively in Table S4."),
]
for p in paras:
    doc.add_paragraph(p)
doc.save(OUT)
print("written", OUT)
